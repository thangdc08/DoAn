# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(13)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_heading_custom(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.bold = True
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(13)
    return h

def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p

def add_normal(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return p

def add_caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.italic = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

# =============================
# MỤC LỤC
# =============================
add_para("MỤC LỤC", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

toc_items = [
    ("TỔNG QUAN", 0),
    ("Đặt vấn đề", 1),
    ("Nhu cầu, công nghệ và mô tả sản phẩm ứng dụng", 1),
    ("Nhu cầu", 2),
    ("Công nghệ sử dụng", 2),
    ("Mô tả ứng dụng", 2),
    ("Phạm vi", 1),
    ("Phạm vi nghiên cứu", 2),
    ("Nội dung thực hiện", 1),
    ("Phương pháp nghiên cứu", 1),
    ("XÂY DỰNG GIẢI PHÁP", 0),
    ("Xây dựng giải pháp công nghệ", 1),
    ("Nền tảng (Platform)", 2),
    ("Ngôn ngữ lập trình", 2),
    ("Thư viện và công cụ phát triển", 2),
    ("Hệ quản trị cơ sở dữ liệu", 2),
    ("Giải pháp hệ thống ứng dụng", 1),
    ("Tổng quan kiến trúc hệ thống", 2),
    ("Cung cấp dịch vụ cho người dùng", 2),
    ("Các chức năng cơ bản", 2),
    ("PHÂN TÍCH & THIẾT KẾ LOGIC HỆ THỐNG", 0),
    ("Phân tích tổng thể hệ thống về chức năng", 1),
    ("Người dùng hệ thống", 2),
    ("Chức năng hệ thống", 2),
    ("Yêu cầu phi chức năng", 2),
    ("Phân tích và thiết kế hệ thống về tổng thể", 1),
    ("Phân tích và thiết kế chức năng hệ thống", 2),
    ("Usecase tổng quát", 3),
    ("Usecase phân rã chức năng", 3),
    ("Phân tích và thiết kế cấu trúc hệ thống", 2),
    ("Xây dựng biểu đồ lớp", 3),
    ("PHÂN TÍCH & THIẾT KẾ CHI TIẾT", 0),
    ("Phân tích chi tiết các chức năng vai trò người dùng", 1),
    ("Usecase đăng nhập", 2),
    ("Usecase đăng xuất", 2),
    ("Usecase xem thông tin cá nhân", 2),
    ("Usecase cập nhật thông tin cá nhân", 2),
    ("Usecase tạo tài khoản", 2),
    ("Phân tích chi tiết các chức năng vai trò chủ sân (Owner)", 2),
    ("Usecase đăng ký sân", 2),
    ("Usecase quản lý sân và sân con", 2),
    ("Usecase quản lý khung giá", 2),
    ("Usecase quản lý đơn đặt sân", 2),
    ("Usecase quản lý doanh thu", 2),
    ("Phân tích chi tiết các chức năng vai trò quản trị viên", 2),
    ("Usecase xem thông tin tài khoản", 3),
    ("Usecase tìm kiếm tài khoản", 3),
    ("Usecase xem danh sách tài khoản", 3),
    ("Usecase phê duyệt sân", 3),
    ("Usecase quản lý báo cáo", 3),
    ("Phân tích chi tiết các chức năng vai trò người dùng (Customer)", 2),
    ("Usecase tìm kiếm sân", 3),
    ("Usecase xem chi tiết sân", 3),
    ("Usecase đặt sân cầu lông", 3),
    ("Usecase thanh toán đơn đặt sân", 3),
    ("Usecase đánh giá sân", 3),
    ("Phân tích chi tiết chức năng cộng đồng (Community)", 2),
    ("Usecase tạo bài đăng kèo giao lưu", 3),
    ("Usecase tham gia kèo giao lưu", 3),
    ("Usecase quản lý người tham gia", 3),
    ("Usecase đánh giá người chơi", 3),
    ("Thiết kế cơ sở dữ liệu", 1),
    ("Mô hình dữ liệu", 2),
    ("Chi tiết các thực thể chính", 2),
    ("Thiết kế giao diện người dùng", 1),
    ("Giao diện người dùng (User)", 2),
    ("Giao diện chủ sân (Owner)", 2),
    ("Giao diện quản trị viên (Admin)", 2),
    ("Giao diện khách vãng lai", 2),
    ("XÂY DỰNG HỆ THỐNG & KIỂM THỬ", 0),
    ("Xây dựng hệ thống", 1),
    ("Kiến trúc triển khai", 2),
    ("Công nghệ triển khai", 2),
    ("Kết quả xây dựng", 2),
    ("Kiểm thử", 1),
    ("Phương pháp kiểm thử", 2),
    ("Kết quả kiểm thử", 2),
    ("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN TƯƠNG LAI", 0),
    ("Kết luận", 1),
    ("Hướng phát triển tương lai", 1),
]

for item, level in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    if level == 0:
        run.bold = True
    elif level == 1:
        pass
    elif level == 2:
        p.paragraph_format.left_indent = Inches(0.3)
    elif level == 3:
        p.paragraph_format.left_indent = Inches(0.6)

doc.add_page_break()

# =============================
# CHƯƠNG 1: TỔNG QUAN
# =============================
add_heading_custom("TỔNG QUAN", 1)

add_heading_custom("Đặt vấn đề", 2)
add_normal("Cầu lông là môn thể thao phổ biến tại Việt Nam với lượng người chơi ngày càng tăng. Thực tế hiện nay, việc đặt sân cầu lông vẫn chủ yếu diễn ra thông qua các phương thức truyền thống như gọi điện thoại, nhắn tin hoặc đến trực tiếp cơ sở để đặt lịch. Phương thức này gây nhiều bất tiện cho cả người chơi và chủ sân: người chơi khó biết được tình trạng sân trống, không thể so sánh giá cả giữa các địa điểm, trong khi chủ sân gặp khó khăn trong việc quản lý lịch đặt, theo dõi doanh thu và tối ưu hóa công suất sân. Đặc biệt, giờ cao điểm buổi tối và cuối tuần thường xuyên xảy ra tình trạng tụt hờn (overbooking) hoặc sân để trống không ai biết.")
add_normal("Bên cạnh đó, cộng đồng người chơi cầu lông còn thiếu một nền tảng kết nối để tìm đối thủ, tổ chức các trận giao lưu (kèo) và đánh giá lẫn nhau. Người chơi thường phải thông qua các nhóm Zalo, Facebook để tìm kèo, quá trình này thiếu minh bạch, khó quản lý và không có cơ chế đánh giá uy tín giữa các người chơi.")
add_normal("Xuất phát từ những bất cập thực tế đó, được sự nhất trí của Giáo viên hướng dẫn và sự chấp nhận của khoa và nhà trường, em đã lựa chọn thực hiện đề tài: \"Nền tảng đặt sân cầu lông trực tuyến tích hợp cộng đồng giao lưu và gợi ý thông minh\". Đề tài hướng đến việc xây dựng một hệ thống thông tin toàn diện, kết nối người chơi với chủ sân thông qua nền tảng web, đồng thời tạo ra một cộng đồng giao lưu lành mạnh cho người yêu thích môn thể thao này.")

add_heading_custom("Nhu cầu, công nghệ và mô tả sản phẩm ứng dụng", 2)

add_heading_custom("Nhu cầu", 3)
add_normal("Xuất phát từ thực trạng thiếu một nền tảng tập trung cho việc đặt sân cầu lông và kết nối cộng đồng người chơi, đề tài xây dựng một giải pháp tích hợp trên nền tảng web responsive. Hệ thống hướng đến ba nhóm người dùng chính: người chơi cầu lông muốn tìm và đặt sân một cách tiện lợi, chủ sân cần quản lý lịch đặt và doanh thu, và quản trị viên cần giám sát hoạt động của toàn hệ thống.")
add_normal("Mô hình hệ thống: Client-Server gồm Web App (React) cho người dùng cuối, Web Admin (React) cho quản trị viên và chủ sân, và Backend API dựa trên kiến trúc microservices (Java Spring Boot) xử lý toàn bộ nghiệp vụ.")
add_normal("Phạm vi chức năng (MVP): Đồ án tập trung phát triển các tính năng cốt lõi gồm: Đặt sân cầu lông trực tuyến với khóa chỗ tự động, Thanh toán trực tuyến (mock), Cộng đồng kèo giao lưu, Đánh giá sân và người chơi, Gợi ý sân và kèo thông minh, Quản lý người dùng và phân quyền.")

add_heading_custom("Công nghệ sử dụng", 3)
add_normal("Hệ thống được xây dựng dựa trên kiến trúc microservices hiện đại, sử dụng bộ công nghệ sau:")
add_normal("Công nghệ phát triển:")
add_normal("Frontend: Sử dụng React kết hợp với Vite và TypeScript cho ứng dụng web responsive. UI được xây dựng với Ant Design cho các trang dashboard và Tailwind CSS cho giao diện công khai. Quản lý trạng thái phía client sử dụng Zustand và TanStack React Query, định tuyến với React Router DOM.")
add_normal("Backend: Vận hành trên nền tảng Java 21 với Spring Boot 3.3.0, sử dụng Spring Cloud Gateway làm cổng vào duy nhất và Netflix Eureka làm service registry cho việc phát hiện dịch vụ. Các microservice giao tiếp thông qua sự kiện với Apache Kafka và gọi REST API với Spring Cloud OpenFeign.")
add_normal("Cơ sở dữ liệu: PostgreSQL 16 với tiện mở rộng PostGIS cho truy vấn không gian địa lý, sử dụng cơ chế đa schema (identity, venue, booking, payment, community) trong cùng một instance. MongoDB 7 lưu trữ dữ liệu thông báo và gợi ý. Redis 7 quản lý OTP, bộ nhớ đệm và khóa chỗ sân.")

add_heading_custom("Mô tả ứng dụng", 3)
add_normal("1.3.1. Hệ thống hoạt động theo mô hình kiến trúc microservices, trong đó mỗi dịch vụ đảm nhận một lĩnh vực nghiệp vụ riêng biệt và giao tiếp với các dịch vụ khác thông qua API Gateway và cơ chế sự kiện Kafka.")
add_normal("a. Phân hệ Người dùng (User): Đây là phân hệ dành cho người chơi cầu lông. Giao diện được thiết kế thân thiện, tập trung vào trải nghiệm tìm kiếm và đặt sân. Người dùng có thể tìm kiếm sân theo vị trí, lịch sử đặt sân, tham gia kèo giao lưu và đánh giá sân cũng như người chơi khác.")
add_normal("b. Phân hệ Chủ sân (Owner): Đây là phân hệ dành cho chủ cơ sở sân cầu lông. Giao diện được tối ưu hóa cho việc quản lý sân, theo dõi đơn đặt sân và doanh thu. Chủ sân có thể đăng ký sân mới, quản lý các sân con, thiết lập khung giá và theo dõi thống kê doanh thu.")
add_normal("c. Phân hệ Quản trị viên (Admin): Đây là phân hệ quản lý toàn hệ thống. Quản trị viên có thể phê duyệt sân mới, quản lý tất cả người dùng, giám sát tất cả đơn đặt sân và xem báo cáo tổng quan của hệ thống.")
add_normal("d. Phân hệ Cộng đồng (Community): Cung cấp chức năng tạo và tham gia các trận kèo giao lưu, đánh giá người chơi và gợi ý kèo phù hợp dựa trên vị trí và trình độ của người dùng.")
add_normal("e. Cơ chế vận hành: Hệ thống sử dụng cơ chế khóa chỗ (slot locking) để tránh tình trạng đặt sân trùng lặp. Khi người dùng chọn khung giờ, hệ thống sẽ khóa các slot đó trong thời gian giới hạn (10-15 phút) để người dùng hoàn tất thanh toán. Nếu hết hạn mà chưa thanh toán, các slot sẽ tự động được mở lại.")

add_heading_custom("Phạm vi", 1)

add_heading_custom("Phạm vi nghiên cứu", 2)
add_normal("- Nghiên cứu các vấn đề lý thuyết, công nghệ xung quanh việc phát triển hệ thống thông tin dựa trên kiến trúc microservices, xử lý sự kiện không đồng bộ và tích hợp các dịch vụ thời gian thực.")
add_normal("- Nghiên cứu các giải pháp đặt sân trực tuyến hiện có trên thị trường để phân tích ưu nhược điểm, từ đó đề xuất kiến trúc và tính năng phù hợp cho hệ thống.")
add_normal("- Áp dụng các kiến thức đã học để phân tích, thiết kế và phát triển hoàn thiện hệ thống đặt sân cầu lông tích hợp cộng đồng giao lưu trên nền tảng web responsive.")
add_normal("- Xây dựng và trình bày báo cáo đồ án tốt nghiệp hoàn chỉnh.")

add_heading_custom("Nội dung thực hiện", 1)
add_normal("Quy trình thực hiện đồ án bao gồm 4 giai đoạn chính:")
add_normal("Nghiên cứu: Tìm hiểu kiến trúc microservices, mô hình giao tiếp sự kiện (event-driven), công nghệ Spring Boot, React, PostgreSQL với PostGIS, Redis, Kafka và các công nghệ liên quan. Nghiên cứu các hệ thống đặt sân hiện có trên thị trường.")
add_normal("Phân tích và Thiết kế: Xây dựng đặc tả nghiệp vụ chi tiết cho từng module, thiết kế kiến trúc tổng thể hệ thống, thiết kế cơ sở dữ liệu với đa schema, thiết kế API contract cho từng microservice và thiết kế giao diện người dùng (UI/UX).")
add_normal("Xây dựng: Tiến hành lập trình Backend (10 microservices), phát triển Frontend (React + Vite + TypeScript), tích hợp các dịch vụ bên thứ ba (Google OAuth, Kafka, Redis), xây dựng hệ thống thông báo và gợi ý thông minh.")
add_normal("Kiểm thử: Thực hiện kiểm thử chức năng toàn hệ thống, kiểm thử tích hợp giữa các microservice, kiểm thử API với Swagger UI và Postman, đánh giá hiệu năng hệ thống.")

add_heading_custom("Phương pháp nghiên cứu", 1)
add_normal("Để hoàn thành đề tài, em thực hiện theo hai phương pháp: nghiên cứu lý thuyết kết hợp nghiên cứu thực nghiệm. Cụ thể như sau:")
add_normal("Nghiên cứu lý thuyết:")
add_normal("- Tìm kiếm và nghiên cứu các tài liệu, sách báo liên quan đến kiến trúc microservices, phát triển ứng dụng web với React và Spring Boot, hệ thống đặt sân trực tuyến và các mô hình cộng đồng trực tuyến.")
add_normal("- Nghiên cứu các công nghệ nền tảng: Spring Boot 3.x, Spring Cloud Gateway, Apache Kafka, PostgreSQL/PostGIS, Redis, React 19, Vite, TypeScript.")
add_normal("Nghiên cứu thực nghiệm:")
add_normal("- Sử dụng các kiến thức đã nghiên cứu để xây dựng đặc tả nghiệp vụ và thiết kế kiến trúc hệ thống.")
add_normal("- Sử dụng các công cụ và framework đã chọn để xây dựng từng microservice và giao diện người dùng.")
add_normal("- Nghiên cứu các sản phẩm, hệ thống tương tự đã có sẵn trên thị trường như SportTrack, CourtCaller, Smash.gg để tìm hiểu về các ưu điểm và chức năng cơ bản cần thiết. Từ đó áp dụng các phát hiện này vào quá trình xây dựng ứng dụng của mình.")
add_normal("- Đánh giá, kiểm thử hệ thống và thu thập phản hồi từ người dùng thử nghiệm để hoàn thiện sản phẩm.")

doc.add_page_break()

# =============================
# CHƯƠNG 2: XÂY DỰNG GIẢI PHÁP
# =============================
add_heading_custom("XÂY DỰNG GIẢI PHÁP", 1)

add_heading_custom("Xây dựng giải pháp công nghệ", 2)
add_normal("Giải pháp công nghệ được đề xuất nhằm đảm bảo hệ thống hoạt động ổn định, bảo mật, có khả năng mở rộng và dễ bảo trì. Hệ thống được xây dựng dựa trên kiến trúc microservices với event-driven communication, đảm bảo tính module hóa cao và khả năng triển khai độc lập cho từng dịch vụ.")

add_heading_custom("Nền tảng (Platform)", 3)
add_normal("Hệ thống được phát triển theo mô hình kiến trúc Client-Server (Khách - Chủ), hoạt động trên các nền tảng hạ tầng và môi trường thực thi sau:")
add_normal("a. Nền tảng ứng dụng web (Web Platform):")
add_normal("Frontend: Ứng dụng web responsive được xây dựng với React và Vite, hoạt động trên mọi trình duyệt hiện đại. Giao diện được thiết kế responsive để hỗ trợ cả máy tính để bàn và thiết bị di động.")
add_normal("Backend: Kiến trúc microservices được triển khai trên Docker containers, mỗi dịch vụ có thể được triển khai và mở rộng độc lập.")
add_normal("b. Nền tảng phát triển ứng dụng (Development Platform):")
add_normal("React + Vite + TypeScript: Giải pháp sử dụng React với Vite làm công cụ build và TypeScript làm ngôn ngữ lập trình chính cho Frontend. Sự kết hợp này đảm bảo tốc độ phát triển nhanh, trải nghiệm phát triển tốt và độ an toàn cao nhờ hệ thống kiểu tĩnh của TypeScript.")
add_normal("Java 21 + Spring Boot 3.x: Nền tảng phát triển Backend sử dụng Java 21 với Spring Boot 3.3.0, cung cấp khả năng xây dựng microservice nhanh chóng với cấu hình tối thiểu.")
add_normal("c. Nền tảng hạ tầng máy chủ và Đám mây (Infrastructure):")
add_normal("Docker và Docker Compose: Toàn bộ hệ thống được đóng gói và triển khai bằng Docker, sử dụng Docker Compose để quản lý orchestration của tất cả các dịch vụ (database, message broker, các microservices).")
add_normal("Apache Kafka: Hệ thống sử dụng Kafka làm cơ chế truyền thông sự kiện (event bus) để giao tiếp không đồng bộ giữa các microservices, đảm bảo tính decoupling và khả năng chịu lỗi cao.")

add_heading_custom("Ngôn ngữ lập trình", 3)
add_normal("Để đảm bảo tính đồng nhất và hiệu quả trong quá trình phát triển hệ thống, đồ án sử dụng hai ngôn ngữ chính sau:")
add_normal("a. Java 21")
add_normal("Vai trò: Là ngôn ngữ lập trình chính cho toàn bộ Backend microservices. Java cung cấp hệ sinh thái Spring Boot mạnh mẽ, phù hợp để xây dựng các dịch vụ doanh nghiệp với độ ổn định và bảo mật cao.")
add_normal("Đặc điểm: Java 21 mang đến các cải tiến về hiệu năng (Virtual Threads), garbage collector mới và các tính năng ngôn ngữ hiện đại, giúp tối ưu hiệu suất xử lý đồng thời của các microservices.")
add_normal("b. TypeScript")
add_normal("Vai trò: Được áp dụng chủ đạo cho phân hệ Frontend với React. TypeScript bổ sung cơ chế kiểu tĩnh, giúp phát hiện lỗi sớm trong quá trình phát triển và tăng khả năng bảo trì mã nguồn.")
add_normal("Đặc điểm: Với hệ thống có hơn 50 trang và hàng trăm thành phần giao diện, TypeScript giúp đảm bảo tính nhất quán về kiểu dữ liệu giữa các dịch vụ Backend và Frontend thông qua các interface và type definition chung.")

add_heading_custom("Thư viện và công cụ phát triển", 3)
add_normal("a. Thư viện Backend (Java):")
add_normal("Spring Boot 3.3.0: Framework chính để xây dựng các microservices, cung cấp auto-configuration và starter dependencies.")
add_normal("Spring Cloud Gateway: Cổng vào duy nhất cho tất cả các yêu cầu từ Frontend, xử lý routing, authentication và rate limiting.")
add_normal("Spring Cloud Netflix Eureka: Service Discovery cho việc đăng ký và tìm kiếm các microservices.")
add_normal("Spring Data JPA: ORM framework để làm việc với PostgreSQL thông qua Hibernate.")
add_normal("Spring Data MongoDB: Làm việc với MongoDB cho Notification Service và Recommendation Service.")
add_normal("Spring Kafka: Thư viện tích hợp Apache Kafka để gửi và nhận sự kiện giữa các dịch vụ.")
add_normal("Spring Security + JWT: Xác thực và phân quyền dựa trên JWT tokens với cơ chế refresh token.")
add_normal("Springdoc OpenAPI: Tự động tạo tài liệu API (Swagger UI) cho mỗi microservice.")
add_normal("PostGIS: Tiện mở rộng của PostgreSQL cho truy vấn địa lý và không gian (tìm sân gần, tìm kèo gần).")
add_normal("b. Thư viện Frontend (TypeScript/React):")
add_normal("React 19.2.6: Thư viện UI chính, sử dụng mô hình component-based và Virtual DOM.")
add_normal("Vite 8.0.12: Công cụ build với tốc độ phát triển nhanh nhờ native ES modules và Hot Module Replacement (HMR).")
add_normal("Ant Design 6.3.7: Bộ thành phần UI chuyên dụng cho các trang dashboard, bảng dữ liệu, biểu mẫu và các thành phần giao diện quản trị.")
add_normal("Tailwind CSS 4.3.0: Framework CSS utility-first cho việc xây dựng giao diện responsive nhanh chóng.")
add_normal("TanStack React Query 5.x: Quản lý trạng thái máy chủ (server state), caching và tối ưu việc fetching dữ liệu từ API.")
add_normal("Zustand: Quản lý trạng thái toàn cục nhẹ cho authentication, chat và community.")
add_normal("Axios: HTTP client với interceptor để tự động gắn JWT token và xử lý refresh token.")
add_normal("React Router DOM 7.x: Định tuyến phía client với các route guard dựa trên vai trò người dùng.")
add_normal("React Hook Form + Zod: Quản lý và xác thực biểu mẫu.")
add_normal("Leaflet + React-Leaflet: Thư viện bản đồ mã nguồn mở cho tính năng tìm kiếm sân và kèo theo vị trí địa lý.")
add_normal("Recharts: Thư viện biểu đồ cho dashboard của Owner và Admin.")
add_normal("DND Kit: Thư viện kéo-thả cho giao diện chọn khung giờ đặt sân.")
add_normal("STOMP.js + SockJS: Giao thức WebSocket cho tính năng chat thời gian thực.")
add_normal("c. Công cụ phát triển và DevOps:")
add_normal("Maven 3.9+: Quản lý phụ thuộc và build cho tất cả các microservices Backend.")
add_normal("Docker và Docker Compose: Đóng gói và triển khai toàn bộ hệ thống với cấu hình reproducible.")
add_normal("Nginx: Web server phục vụ Frontend và reverse proxy cho API Gateway.")

add_heading_custom("Hệ quản trị cơ sở dữ liệu", 3)
add_normal("Hệ thống sử dụng chiến lược đa cơ sở dữ liệu (polyglot persistence), mỗi loại database được chọn phù hợp với đặc thù nghiệp vụ của từng module:")
add_normal("a. PostgreSQL 16 với PostGIS:")
add_normal("Vai trò: Cơ sở dữ liệu chính, lưu trữ toàn bộ dữ liệu nghiệp vụ cốt lõi của hệ thống thông qua cơ chế đa schema.")
add_normal("Cấu trúc schema:")
add_normal("- Schema identity: Quản lý tài khoản người dùng, vai trò, refresh tokens, hồ sơ cá nhân và cài đặt người dùng.")
add_normal("- Schema venue: Quản lý thông tin sân, sân con, khung giờ hoạt động, quy tắc giá, hình ảnh sân và đánh giá sân.")
add_normal("- Schema booking: Quản lý đơn đặt sân, chi tiết đơn và khóa chỗ (slot lock) để đảm bảo không đặt trùng.")
add_normal("- Schema payment: Quản lý giao dịch thanh toán.")
add_normal("- Schema community: Quản lý bài đăng kèo giao lưu, người tham gia, đánh giá người chơi và các tương tác xã hội.")
add_normal("PostGIS: Tiện mở rộng không gian cho phép thực hiện các truy vấn tìm kiếm theo vị trí địa lý (\"tìm sân gần tôi\", \"tìm kèo gần tôi\") một cách hiệu quả thông qua các toán tử không gian như ST_DWithin.")
add_normal("b. MongoDB 7:")
add_normal("Vai trò: Lưu trữ dữ liệu phi cấu trúc và bán cấu trúc, phù hợp với mô hình dữ liệu linh hoạt của các module thông báo và gợi ý.")
add_normal("Cấu trúc database:")
add_normal("- notification_db: Lưu trữ thông báo của người dùng với các chỉ mục được tối ưu để truy vấn theo người nhận và thời gian tạo.")
add_normal("- recommendation_db: Lưu trữ lịch sử hoạt động của người dùng để phục vụ cho thuật toán gợi ý cá nhân hóa.")
add_normal("c. Redis 7:")
add_normal("Vai trò: Hệ thống bộ nhớ đệm (in-memory) đa mục đích, cung cấp hiệu năng cao cho các thao tác cần tốc độ phản hồi nhanh.")
add_normal("Ứng dụng:")
add_normal("- Lưu trữ OTP (One-Time Password) cho chức năng khôi phục mật khẩu với thời gian sống (TTL) 10 phút.")
add_normal("- Bộ nhớ đệm kết quả gợi ý sân (10 phút) và gợi ý kèo (5 phút) để giảm tải cho các dịch vụ nền.")
add_normal("- Quản lý khóa chỗ sân (slot lock) với TTL tự động hết hạn.")
add_normal("d. Apache Kafka 3.7.0:")
add_normal("Vai trò: Cơ chế truyền thông sự kiện (event bus) trung tâm, cho phép các microservices giao tiếp không đồng bộ thông qua các topic.")
add_normal("Chức năng: Đảm bảo tính decoupling giữa các dịch vụ, cho phép mở rộng dễ dàng bằng cách thêm consumer mới mà không cần thay đổi producer.")

doc.add_page_break()

add_heading_custom("Giải pháp hệ thống ứng dụng", 2)
add_normal("Hệ thống được thiết kế theo kiến trúc microservices, trong đó mỗi dịch vụ đảm nhận một lĩnh vực nghiệp vụ riêng biệt và giao tiếp với các dịch vụ khác thông qua API Gateway và cơ chế sự kiện Kafka.")

add_heading_custom("Tổng quan kiến trúc hệ thống", 3)
add_normal("Hệ thống được triển khai theo kiến trúc microservices với các thành phần chính sau:")
add_normal("API Gateway (Spring Cloud Gateway - Port 8080): Là điểm vào duy nhất cho tất cả các yêu cầu từ Frontend. Thực hiện routing đến các microservice tương ứng, xác thực JWT token và áp dụng các chính sách cross-cutting như rate limiting.")
add_normal("Service Registry (Netflix Eureka - Port 8761): Cung cấp cơ chế phát hiện dịch vụ (service discovery), cho phép các microservice đăng ký và tìm kiếm lẫn nhau một cách động.")
add_normal("Config Server (Spring Cloud Config - Port 8888): Quản lý cấu hình tập trung cho tất cả các microservices, hỗ trợ thay đổi cấu động mà không cần restart dịch vụ.")
add_normal("Identity Service (Port 8081): Quản lý xác thực người dùng, phân quyền, hồ sơ cá nhân và khôi phục mật khẩu.")
add_normal("Venue Service (Port 8082): Quản lý thông tin sân, sân con, khung giờ, giá cả và đánh giá sân. Tích hợp PostGIS cho tìm kiếm theo vị trí.")
add_normal("Booking Service (Port 8083): Quản lý quy trình đặt sân, bao gồm cơ chế khóa chỗ (slot lock) để ngăn đặt trùng và quản lý vòng đời đơn đặt sân.")
add_normal("Payment Service (Port 8084): Xử lý thanh toán cho các đơn đặt sân, tích hợp cổng thanh toán (mock payment và VNPay).")
add_normal("Community Service (Port 8085): Quản lý cộng đồng người chơi, bao gồm bài đăng kèo giao lưu, quản lý người tham gia, đánh giá người chơi và tìm kiếm kèo theo vị trí.")
add_normal("Notification Service (Port 8086): Quản lý thông báo cho người dùng dựa trên các sự kiện từ các microservice khác.")
add_normal("Recommendation Service (Port 8087): Cung cấp gợi ý sân và kèo giao lưu thông minh dựa trên hành vi và vị trí của người dùng, sử dụng Redis làm bộ nhớ đệm.")
add_normal("Chat Service (Port 8686): Cung cấp tính năng nhắn tin thời gian thực giữa người chơi thông qua WebSocket.")

add_heading_custom("Cung cấp dịch vụ cho người dùng", 3)
add_normal("Hệ thống được thiết kế theo mô hình đa vai trò, cung cấp các nhóm dịch vụ chuyên biệt phù hợp với từng loại người dùng:")
add_normal("a. Nhóm dịch vụ dành cho Người chơi (User):")
add_normal("Tìm kiếm và đặt sân: Cho phép người chơi tìm kiếm sân theo vị trí địa lý, xem chi tiết thông tin sân (giá, loại sân, đánh giá), chọn khung giờ trống, khóa chỗ và hoàn tất đặt sân.")
add_normal("Thanh toán trực tuyến: Thực hiện thanh toán cho đơn đặt sân thông qua cổng thanh toán tích hợp.")
add_normal("Cộng đồng kèo giao lưu: Tạo và tham gia các trận kèo giao lưu, theo dõi trạng thái tham gia và tương tác với cộng đồng người chơi.")
add_normal("Đánh giá và xếp hạng: Đánh giá sân sau khi sử dụng và đánh giá người chơi sau khi tham gia kèo, xây dựng hệ thống uy tín cá nhân.")
add_normal("Thông báo thời gian thực: Nhận thông báo về đơn đặt sân, lời mời tham gia kèo, kết quả đánh giá và các hoạt động liên quan.")
add_normal("Gợi ý thông minh: Nhận gợi ý sân phù hợp dựa trên lịch sử đặt sân, vị trí và sở thích cá nhân.")
add_normal("b. Nhóm dịch vụ dành cho Chủ sân (Owner):")
add_normal("Quản lý sân: Đăng ký sân mới, cập nhật thông tin sân, quản lý các sân con, thiết lập khung giá và giờ hoạt động.")
add_normal("Quản lý đơn đặt sân: Xem danh sách đơn đặt, xác nhận hoặc hủy đơn, theo dõi trạng thái từng đơn.")
add_normal("Quản lý doanh thu: Xem thống kê doanh thu theo ngày, tuần, tháng; theo dõi tình trạng công suất sân.")
add_normal("Phê duyệt người tham gia: Duyệt hoặc từ chối yêu cầu tham gia kèo (đối với các kèo yêu cầu phê duyệt).")
add_normal("c. Nhóm dịch vụ dành cho Quản trị viên (Admin):")
add_normal("Quản lý người dùng: Xem danh sách tất cả người dùng, tìm kiếm, khóa/mở khóa tài khoản, thay đổi vai trò, xóa tài khoản.")
add_normal("Phê duyệt sân: Xem danh sách sân chờ phê duyệt, phê duyệt hoặc từ chối đăng ký sân mới.")
add_normal("Giám sát hệ thống: Xem tất cả đơn đặt sân trên hệ thống, hủy đơn nếu cần, xem báo cáo tổng quan.")
add_normal("Quản lý báo cáo: Xem thống kê tổng quan về hoạt động của hệ thống.")

add_heading_custom("Các chức năng cơ bản", 3)
add_normal("Dựa trên phân tích hệ thống, các chức năng được chia thành các nhóm chính sau:")
add_normal("a. Nhóm chức năng Hệ thống và Tài khoản (Chung):")
add_normal("Đăng ký/Đăng nhập (Authentication): Hệ thống hỗ trợ đăng ký tài khoản mới, đăng nhập với email và mật khẩu. Xác thực được thực hiện thông qua JWT access token và refresh token. Hệ thống hỗ trợ phân quyền tự động dựa trên vai trò (Người chơi/Chủ sân/Quản trị viên).")
add_normal("Quản lý hồ sơ cá nhân (Profile Management): Người dùng có thể cập nhật thông tin cá nhân bao gồm tên, giới tính, trình độ chơi, mục tiêu và khu vực ưa thích. Chủ sân có thể cập nhật thông tin liên hệ và mô tả sân.")
add_normal("Khôi phục mật khẩu: Chức năng khôi phục mật khẩu thông qua OTP (mã xác thực 6 chữ số) được lưu trữ trong Redis với thời gian sống 10 phút.")
add_normal("Đăng xuất: Kết thúc phiên làm việc an toàn bằng cách thu hồi refresh token trong cơ sở dữ liệu.")
add_normal("b. Nhóm chức năng nghiệp vụ Người chơi:")
add_normal("Tìm kiếm và xem sân: Tìm kiếm sân theo khu vực, xem thông tin chi tiết sân (loại sân, giá, đánh giá, hình ảnh), xem lịch sử đánh giá.")
add_normal("Đặt sân cầu lông: Chọn sân, chọn ngày và khung giờ trống, xác nhận đặt sân. Hệ thống tự động khóa chỗ trong thời gian giới hạn để tránh đặt trùng.")
add_normal("Thanh toán: Thanh toán đơn đặt sân qua cổng thanh toán tích hợp (VNPay hoặc thanh toán thử nghiệm).")
add_normal("Xem lịch sử đặt sân: Theo dõi các đơn đặt sân đã thực hiện, theo dõi trạng thái và lịch sử thanh toán.")
add_normal("Đánh giá sân: Đánh giá sân sau khi đã thanh toán đơn đặt sân thành công. Đánh giá bao gồm số sao (1-5) và nhận xét.")
add_normal("Tham gia cộng đồng: Tạo bài đăng kèo giao lưu, tham gia kèo của người khác, theo dõi trạng thái tham gia, đánh giá người chơi sau khi tham gia kèo.")
add_normal("c. Nhóm chức năng nghiệp vụ Chủ sân:")
add_normal("Đăng ký sân: Đăng ký thông tin sân mới, chờ quản trị viên phê duyệt trước khi hiển thị công khai.")
add_normal("Quản lý sân con: Thêm, sửa, xóa các sân con thuộc sân, cấu hình loại sân (đơn/đôi) và trạng thái hoạt động.")
add_normal("Quản lý khung giờ và giá: Thiết lập khung giờ hoạt động theo ngày trong tuần và quy tắc giá linh hoạt theo thời gian.")
add_normal("Xử lý đơn đặt sân: Xem và xử lý các đơn đặt sân đến, theo dõi doanh thu theo thời gian thực.")
add_normal("Quản lý doanh thu: Xem thống kê doanh thu, biểu đồ phân tích doanh thu theo các mốc thời gian.")
add_normal("d. Nhóm chức năng nghiệp vụ Quản trị viên:")
add_normal("Quản lý người dùng: Xem danh sách, tìm kiếm, khóa/mở khóa tài khoản, thay đổi vai trò và xóa người dùng.")
add_normal("Phê duyệt sân: Xem danh sách sân chờ phê duyệt, phê duyệt hoặc từ chối sân với lý do.")
add_normal("Giám sát hệ thống: Xem tất cả đơn đặt sân trên hệ thống, hủy đơn nếu cần thiết, xem báo cáo tổng quan.")
add_normal("Quản lý báo cáo: Xem thống kê tổng quan về người dùng, đơn đặt sân và doanh thu của hệ thống.")

doc.add_page_break()

# =============================
# CHƯƠNG 3: PHÂN TÍCH & THIẾT KẾ LOGIC HỆ THỐNG
# =============================
add_heading_custom("PHÂN TÍCH & THIẾT KẾ LOGIC HỆ THỐNG", 1)

add_heading_custom("Phân tích tổng thể hệ thống về chức năng", 2)
add_normal("Phân tích tổng thể hệ thống được thực hiện nhằm xác định rõ các chức năng, tác nhân và luồng tương tác giữa các thành phần, từ đó làm cơ sở cho việc thiết kế chi tiết từng module.")
add_normal("Mục tiêu phân tích: Xây dựng hệ thống đặt sân cầu lông trực tuyến tích hợp cộng đồng giao lưu và gợi ý thông minh, giúp kết nối người chơi với chủ sân một cách hiệu quả, đồng thời tạo ra một cộng đồng giao lưu lành mạnh cho người yêu thích môn thể thao cầu lông. Quản lý hiệu quả và bảo mật thông tin người dùng, lịch sử đặt sân cũng như các giao dịch thanh toán, đảm bảo sự vận hành ổn định và tin cậy của toàn hệ thống.")

add_heading_custom("Người dùng hệ thống", 3)
add_normal("Từ các thông tin phân tích nghiệp vụ, các tác nhân (actors) tham gia vào hệ thống được xác định như sau:")
add_normal("Khách vãng lai (Guest): Người dùng chưa đăng nhập, chỉ có thể xem thông tin sân công khai và duyệt các kèo giao lưu đang mở. Không có quyền đặt sân hay tương tác sâu.")
add_normal("Người chơi (User): Người dùng đã đăng ký tài khoản và đăng nhập. Có quyền tìm kiếm sân, đặt sân, thanh toán, tham gia kèo giao lưu, đánh giá sân và người chơi, nhận thông báo.")
add_normal("Chủ sân (Owner): Người dùng có vai trò quản lý cơ sở sân cầu lông. Có quyền đăng ký sân mới, quản lý thông tin sân và sân con, thiết lập giá cả, xử lý đơn đặt sân và xem thống kê doanh thu.")
add_normal("Quản trị viên (Admin): Người dùng có quyền quản lý toàn hệ thống. Có quyền phê duyệt sân mới, quản lý tất cả người dùng, giám sát tất cả đơn đặt sân, xem báo cáo tổng quan.")

add_heading_custom("Chức năng hệ thống", 3)
add_normal("Tác nhân chung: Các vai trò quản trị viên, chủ sân và người chơi cùng thực hiện các chức năng cơ bản sau để tương tác với hệ thống:")
add_normal("Đăng nhập: Xác thực người dùng và cấp quyền truy cập theo vai trò.")
add_normal("Đăng xuất: Kết thúc phiên làm việc an toàn.")
add_normal("Quản lý tài khoản cá nhân: Cập nhật thông tin cá nhân và đổi mật khẩu.")
add_normal("Người chơi (User): Vai trò cho người đặt sân và tham gia cộng đồng cầu lông với các chức năng cốt lõi:")
add_normal("- Tìm kiếm và xem thông tin sân cầu lông theo vị trí và bộ lọc")
add_normal("- Đặt sân cầu lông trực tuyến với cơ chế khóa chỗ tự động")
add_normal("- Thanh toán đơn đặt sân")
add_normal("- Xem lịch sử đặt sân và thông báo cá nhân")
add_normal("- Tạo và tham gia kèo giao lưu cầu lông")
add_normal("- Đánh giá sân và người chơi sau khi sử dụng dịch vụ")
add_normal("- Nhận gợi ý sân và kèo phù hợp từ hệ thống")
add_normal("Chủ sân (Owner): Vai trò cho chủ cơ sở sân cầu lông với các chức năng quản lý:")
add_normal("- Đăng ký và quản lý thông tin sân")
add_normal("- Quản lý sân con, khung giờ hoạt động và quy tắc giá")
add_normal("- Xem và xử lý đơn đặt sân đến")
add_normal("- Xem thống kê doanh thu và công suất sân")
add_normal("Quản trị viên (Admin): Vai trò quản lý toàn hệ thống với các chức năng:")
add_normal("- Phê duyệt sân mới đăng ký")
add_normal("- Quản lý tất cả người dùng (xem, tìm kiếm, khóa/mở khóa, thay đổi vai trò, xóa)")
add_normal("- Giám sát tất cả đơn đặt sân và hủy đơn nếu cần")
add_normal("- Xem báo cáo tổng quan về hoạt động hệ thống")

add_heading_custom("Yêu cầu phi chức năng", 3)
add_normal("Bảo mật: Hệ thống cần đảm bảo dữ liệu người dùng và giao dịch thanh toán không bị truy cập trái phép. Mọi yêu cầu API phải được xác thực thông qua JWT token. Thông tin nhạy cảm như mật khẩu được mã hóa bằng BCrypt. Các endpoint chia sẻ giữa các microservice được bảo vệ bằng cơ chế xác thực nội bộ.")
add_normal("Hiệu suất: Hệ thống phải đáp ứng yêu cầu thời gian phản hồi thấp (dưới 500ms) cho các thao tác thông thường. Cơ chế khóa chỗ sân phải xử lý được đồng thời nhiều yêu cầu mà không xảy ra xung đột. Bộ nhớ đệm Redis được sử dụng để giảm tải cho các truy vấn lặp lại.")
add_normal("Tính ổn định: Hệ thống cần hoạt động ổn định với kiến trúc microservices, cho phép một dịch vụ gặp sự cố không ảnh hưởng đến toàn bộ hệ thống. Cơ chế retry và circuit breaker được áp dụng cho các gọi dịch vụ liên quan. Cơ chế sự kiện Kafka đảm bảo các thông báo không bị mất ngay cả khi dịch vụ consumer tạm thời không khả dụng.")
add_normal("Khả năng mở rộng: Kiến trúc microservices cho phép mở rộng từng dịch vụ độc lập dựa trên tải thực tế. Cơ sở dữ liệu được thiết kế theo hướng phân tách schema, hỗ trợ việc tách database instance khi cần.")
add_normal("Khả năng bảo trì: Mã nguồn được tổ chức theo cấu trúc module rõ ràng, mỗi microservice có vòng đời riêng. Tài liệu API tự động được tạo bởi Swagger cho mỗi dịch vụ.")

doc.add_page_break()

add_heading_custom("Phân tích và thiết kế hệ thống về tổng thể", 2)
add_normal("Thiết kế tổng thể hệ thống được xây dựng dựa trên kiến trúc microservices với event-driven communication, đảm bảo tính module hóa cao, khả năng mở rộng và dễ bảo trì.")

add_heading_custom("Phân tích và thiết kế chức năng hệ thống", 3)

add_heading_custom("Usecase tổng quát", 4)
add_normal("Căn cứ vào các thông tin phân tích về người dùng và các chức năng cơ bản của hệ thống, biểu đồ usecase tổng quát đã được xây dựng để thể hiện toàn bộ các chức năng mà hệ thống cung cấp cho các tác nhân tham gia.")
add_normal("Các usecase chính trong hệ thống bao gồm: Đăng ký tài khoản, Đăng nhập, Đăng xuất, Quản lý hồ sơ cá nhân, Tìm kiếm sân, Xem chi tiết sân, Đặt sân, Thanh toán đơn đặt sân, Đánh giá sân, Tạo kèo giao lưu, Tham gia kèo giao lưu, Đánh giá người chơi, Quản lý thông báo, Quản lý sân (Owner), Quản lý đơn đặt (Owner), Quản lý doanh thu (Owner), Quản lý người dùng (Admin), Phê duyệt sân (Admin), Quản lý báo cáo (Admin).")

add_heading_custom("Usecase phân rã chức năng", 4)
add_normal("Từ biểu đồ Usecase tổng quát, các biểu đồ usecase phân rã được phân tách theo từng tác nhân, trong đó các usecase cơ bản được tách và làm rõ chi tiết:")
add_normal("Phân rã tác nhân Người chơi (User): Bao gồm đăng ký tài khoản, đăng nhập, đăng xuất, quản lý hồ sơ cá nhân, tìm kiếm sân, đặt sân, thanh toán, xem lịch sử đặt sân, đánh giá sân, tạo kèo giao lưu, tham gia kèo, đánh giá người chơi và quản lý thông báo.")
add_normal("Phân rã tác nhân Chủ sân (Owner): Bao gồm đăng nhập, quản lý hồ sơ, đăng ký sân, quản lý sân con, quản lý khung giờ và giá, xử lý đơn đặt sân, xem danh sách đơn đặt và quản lý doanh thu.")
add_normal("Phân rã tác nhân Quản trị viên (Admin): Bao gồm đăng nhập, quản lý tài khoản người dùng (xem, tìm kiếm, khóa/mở khóa, thay đổi vai trò, xóa), phê duyệt sân mới, xem tất cả đơn đặt sân, hủy đơn và quản lý báo cáo tổng quan.")

add_heading_custom("Phân tích và thiết kế cấu trúc hệ thống", 3)

add_heading_custom("Xây dựng biểu đồ lớp", 4)
add_normal("Căn cứ vào các thông tin phân tích usecase, các lớp đối tượng được phát hiện và xác định thuộc tính, phương thức ban đầu. Các lớp chính trong hệ thống bao gồm:")
add_normal("Lớp xác thực và người dùng (Auth & User): User, Role, RefreshToken, UserProfile, PasswordResetOtp. Lớp này quản lý toàn bộ thông tin đăng nhập, phân quyền và hồ sơ cá nhân của người dùng.")
add_normal("Lớp sân (Venue): Venue, Court, CourtSlot, PriceRule, VenueImage, VenueRating. Lớp này quản lý toàn bộ thông tin về cơ sở sân cầu lông, các sân con, khung giờ hoạt động, quy tắc giá và đánh giá.")
add_normal("Lớp đặt sân (Booking): Booking, BookingItem, SlotLock. Lớp này quản lý quy trình đặt sân với cơ chế khóa chỗ để đảm bảo không xảy ra tình trạng đặt trùng.")
add_normal("Lớp thanh toán (Payment): PaymentTransaction. Lớp này quản lý các giao dịch thanh toán liên quan đến đơn đặt sân.")
add_normal("Lớp cộng đồng (Community): MatchPost, Participant, PlayerRating, Like, Comment, Report. Lớp này quản lý các bài đăng kèo giao lưu, người tham gia và hệ thống đánh giá người chơi.")
add_normal("Lớp thông báo (Notification): Notification. Lớp này quản lý thông báo được gửi đến người dùng dựa trên các sự kiện trong hệ thống.")
add_normal("Lớp gợi ý (Recommendation): UserActivity. Lớp này lưu trữ lịch sử hoạt động của người dùng để phục vụ cho thuật toán gợi ý.")

doc.add_page_break()

# =============================
# CHƯƠNG 4: PHÂN TÍCH & THIẾT KẾ CHI TIẾT
# =============================
add_heading_custom("PHÂN TÍCH & THIẾT KẾ CHI TIẾT", 1)

add_heading_custom("Phân tích chi tiết các chức năng vai trò người dùng", 2)

add_heading_custom("Usecase đăng nhập", 3)
add_normal("Chức năng đăng nhập là một trong những chức năng nền tảng quan trọng nhất, đóng vai trò là cổng vào cho tất cả các tác nhân vào hệ thống.")
add_normal("Luồng nghiệp vụ: Người dùng nhập email và mật khẩu vào giao diện đăng nhập. Hệ thống kiểm tra tính hợp lệ của thông tin đăng nhập, xác minh trạng thái tài khoản (ACTIVE/LOCKED). Nếu hợp lệ, hệ thống tạo cặp JWT token (access token và refresh token), lưu refresh token đã được mã hóa vào cơ sở dữ liệu và trả về cho người dùng. Access token được sử dụng cho các yêu cầu API tiếp theo, còn refresh token được dùng để làm mới access token khi hết hạn.")
add_normal("Xử lý lỗi: Nếu thông tin đăng nhập không đúng, hệ thống trả về thông báo lỗi rõ ràng. Nếu tài khoản bị khóa, hệ thống thông báo và hướng dẫn liên hệ quản trị viên.")
add_normal("Bảo mật: Mật khẩu được lưu trữ dưới dạng băm (hash) bằng BCrypt, không lưu dạng plain text. Refresh token được lưu với thời hạn và có cơ chế xoay vòng (rotation) khi sử dụng.")

add_heading_custom("Usecase đăng xuất", 3)
add_normal("Chức năng đăng xuất đảm bảo kết thúc phiên làm việc an toàn bằng cách thu hồi refresh token.")
add_normal("Luồng nghiệp vụ: Người dùng chọn đăng xuất. Hệ thống thu hồi (revoke) refresh token hiện tại trong cơ sở dữ liệu, đánh dấu token đó không còn hợp lệ. Client xóa toàn bộ thông tin xác thực khỏi bộ nhớ cục bộ.")

add_heading_custom("Usecase xem thông tin cá nhân", 3)
add_normal("Người dùng có thể xem thông tin hồ sơ cá nhân của mình bao gồm tên, email, số điện thoại, giới tính, trình độ chơi cầu lông, mục tiêu tập luyện và khu vực ưa thích.")
add_normal("Luồng nghiệp vụ: Người dùng truy cập trang hồ sơ cá nhân. Frontend gửi yêu cầu với access token đến API Gateway. Identity Service xác thực token, truy vấn thông tin người dùng và hồ sơ từ database, trả về dữ liệu đầy đủ cho Frontend hiển thị.")

add_heading_custom("Usecase cập nhật thông tin cá nhân", 3)
add_normal("Người dùng có thể cập nhật thông tin hồ sơ cá nhân như tên, giới tính, trình độ chơi, mục tiêu và khu vực ưa thích.")
add_normal("Luồng nghiệp vụ: Người dùng chỉnh sửa thông tin trên giao diện và gửi yêu cầu cập nhật. Hệ thống kiểm tra xác thực, xác thực dữ liệu đầu vào, cập nhật thông tin vào cơ sở dữ liệu và trả về kết quả thành công.")

add_heading_custom("Usecase tạo tài khoản", 3)
add_normal("Chức năng đăng ký tài khoản cho phép người dùng mới tạo tài khoản trong hệ thống.")
add_normal("Luồng nghiệp vụ: Người dùng nhập thông tin đăng ký (email, mật khẩu, tên đầy đủ). Hệ thống kiểm tra email chưa tồn tại, mã hóa mật khẩu bằng BCrypt, tạo tài khoản mới với vai trò mặc định là USER và trả về kết quả thành công. Người dùng có thể đăng nhập ngay sau khi đăng ký.")
add_normal("Đăng ký chủ sân: Người dùng có thể đăng ký với vai trò OWNER, trong đó thông tin cơ sở sân được thu thập thêm và chuyển vào trạng thái chờ phê duyệt.")

add_heading_custom("Phân tích chi tiết các chức năng vai trò chủ sân (Owner)", 2)

add_heading_custom("Usecase đăng ký sân", 3)
add_normal("Chức năng đăng ký sân cho phép người dùng có vai trò OWNER tạo mới một cơ sở sân cầu lông trong hệ thống.")
add_normal("Luồng nghiệp vụ: Chủ sân điền thông tin sân bao gồm tên sân, địa chỉ, thành phố, quận/huyện, tọa độ địa lý (latitude, longitude), số điện thoại liên hệ và mô tả. Hệ thống lưu thông tin sân với trạng thái PENDING_APPROVAL, gửi thông báo cho quản trị viên để phê duyệt. Sau khi được phê duyệt, sân sẽ hiển thị công khai cho người dùng tìm kiếm.")
add_normal("Quy tắc nghiệp vụ: Một chủ sân có thể quản lý nhiều sân. Sân mới đăng ký chỉ hiển thị cho chủ sân và quản trị viên cho đến khi được phê duyệt. Quản trị viên có quyền từ chối sân với lý do rõ ràng.")

add_heading_custom("Usecase quản lý sân và sân con", 3)
add_normal("Chức năng quản lý sân và sân con cho phép chủ sân cấu hình chi tiết cơ sở của mình.")
add_normal("Quản lý thông tin sân: Cập nhật tên, địa chỉ, mô tả, hình ảnh và tiện ích của sân.")
add_normal("Quản lý sân con: Thêm, sửa đổi và xóa các sân con thuộc sân. Mỗi sân con có tên riêng, loại sân (đơn/đôi), trạng thái hoạt động. Chủ sân có thể tạm tắt một sân con để bảo trì hoặc ngừng kinh doanh mà không ảnh hưởng đến các sân con khác.")
add_normal("Quản lý hình ảnh: Tải lên và sắp xếp thứ tự hiển thị hình ảnh sân để thu hút người chơi.")

add_heading_custom("Usecase quản lý khung giá", 3)
add_normal("Chức năng quản lý khung giá cho phép chủ sân thiết lập giá linh hoạt theo thời gian.")
add_normal("Khung giờ hoạt động: Chủ sân thiết lập các khung giờ hoạt động trong tuần (ví dụ: thứ 2 đến chủ nhật, 6:00 đến 23:00). Các slot nằm ngoài khung giờ này sẽ không thể đặt.")
add_normal("Quy tắc giá: Chủ sân có thể tạo nhiều quy tắc giá khác nhau, ví dụ: giá giờ thường, giá giờ cao điểm (chiều tối, cuối tuần), giá khuyến mãi theo thời gian. Mỗi quy tắc giá được áp dụng cho một khoảng thời gian cụ thể trong ngày và áp dụng cho sân con được chọn.")
add_normal("Cơ chế: Khi người dùng đặt sân, hệ thống tự động tính giá dựa trên quy tắc giá phù hợp với khung giờ đã chọn.")

add_heading_custom("Usecase quản lý đơn đặt sân", 3)
add_normal("Chức năng quản lý đơn đặt sân cho phép chủ sân theo dõi và xử lý các đơn đặt sân từ người chơi.")
add_normal("Danh sách đơn đặt: Chủ sân xem danh sách tất cả đơn đặt sân của các sân mình quản lý, có thể lọc theo trạng thái (chờ xác nhận, đã xác nhận, đã thanh toán, đã hủy) và theo khoảng thời gian.")
add_normal("Chi tiết đơn đặt: Xem thông tin chi tiết đơn đặt bao gồm người đặt, sân, khung giờ, tổng tiền và trạng thái thanh toán.")
add_normal("Trạng thái đơn đặt: Đơn đặt sân có các trạng thái sau: PENDING (chờ thanh toán), PAID (đã thanh toán), FAILED (thanh toán thất bại), EXPIRED (đã hết hạn), CANCELLED_BY_ADMIN (bị quản trị viên hủy).")

add_heading_custom("Usecase quản lý doanh thu", 3)
add_normal("Chức năng quản lý doanh thu cung cấp cho chủ sân cái nhìn tổng quan về tình hình kinh doanh.")
add_normal("Thống kê doanh thu: Hiển thị tổng doanh thu theo ngày, tuần, tháng và năm. Có thể xem doanh thu chi tiết theo từng sân con.")
add_normal("Biểu đồ phân tích: Trực quan hóa xu hướng doanh thu qua các biểu đồ, giúp chủ sân phân tích và đưa ra quyết định kinh doanh.")
add_normal("Công suất sân: Theo dõi tỷ lệ sử dụng sân theo thời gian, xác định các khung giờ cao điểm và thấp điểm để tối ưu giá cả.")

add_heading_custom("Phân tích chi tiết các chức năng vai trò quản trị viên", 2)

add_heading_custom("Usecase xem thông tin tài khoản", 3)
add_normal("Quản trị viên có thể xem thông tin chi tiết của bất kỳ người dùng nào trong hệ thống, bao gồm thông tin cá nhân, vai trò, trạng thái tài khoản và lịch sử hoạt động.")

add_heading_custom("Usecase tìm kiếm tài khoản", 3)
add_normal("Quản trị viên có thể tìm kiếm người dùng theo email, tên hoặc số điện thoại. Kết quả tìm kiếm được hiển thị dưới dạng danh sách có phân trang, cho phép xử lý hàng loạt (khóa/mở khóa, thay đổi vai trò).")

add_heading_custom("Usecase xem danh sách tài khoản", 3)
add_normal("Quản trị viên xem danh sách tất cả người dùng trong hệ thống với khả năng lọc theo vai trò (USER/OWNER/ADMIN) và trạng thái (ACTIVE/LOCKED). Danh sách được hiển thị dưới dạng bảng có phân trang, mỗi trang hiển thị tối đa 20 người dùng.")

add_heading_custom("Usecase phê duyệt sân", 3)
add_normal("Quản trị viên xem danh sách các sân mới đăng ký đang chờ phê duyệt. Đối với mỗi sân, quản trị viên xem thông tin chi tiết và quyết định phê duyệt hoặc từ chối. Nếu từ chối, quản trị viên phải cung cấp lý do. Sân được phê duyệt sẽ hiển thị công khai cho người dùng tìm kiếm.")

add_heading_custom("Usecase quản lý báo cáo", 3)
add_normal("Quản trị viên xem các báo cáo tổng quan về hoạt động của hệ thống, bao gồm số lượng người dùng đăng ký theo thời gian, số lượng đơn đặt sân, doanh thu tổng hợp và các chỉ số quan trọng khác. Báo cáo được trực quan hóa thông qua biểu đồ.")

add_heading_custom("Phân tích chi tiết các chức năng vai trò người chơi (Customer)", 2)

add_heading_custom("Usecase tìm kiếm sân", 3)
add_normal("Người chơi có thể tìm kiếm sân cầu lông theo nhiều tiêu chí khác nhau để tìm được sân phù hợp với nhu cầu.")
add_normal("Tìm kiếm theo vị trí: Sử dụng PostGIS, hệ thống tìm các sân nằm trong bán kính nhất định từ vị trí hiện tại của người dùng, sắp xếp theo khoảng cách.")
add_normal("Lọc theo tiêu chí: Người dùng có thể lọc sân theo thành phố, quận/huyện, khoảng giá, tiện ích (đổ xe, có nhà vệ sinh, có nước uống...) và từ khóa tên sân.")
add_normal("Kết quả hiển thị: Danh sách sân được hiển thị dưới dạng thẻ (card) với hình ảnh, tên, địa chỉ, giá từ và đánh giá trung bình. Người dùng có thể chuyển sang chế độ xem bản đồ để khám phá sân trực quan.")

add_heading_custom("Usecase xem chi tiết sân", 3)
add_normal("Người chơi có thể xem thông tin chi tiết của một sân bao gồm thông tin cơ bản, các sân con, khung giờ hoạt động, bảng giá, hình ảnh và đánh giá từ người chơi khác.")
add_normal("Thông tin sân: Tên sân, địa chỉ, số điện thoại liên hệ, mô tả, tiện ích có sẵn và tọa độ trên bản đồ.")
add_normal("Sân con: Danh sách các sân con với loại sân (đơn/đôi), trạng thái hoạt động và giá theo giờ.")
add_normal("Đánh giá: Danh sách các đánh giá từ người chơi đã sử dụng sân, bao gồm số sao và nhận xét.")

add_heading_custom("Usecase đặt sân cầu lông", 3)
add_normal("Quy trình đặt sân là chức năng cốt lõi của hệ thống, được thiết kế với cơ chế khóa chỗ để đảm bảo không xảy ra tình trạng đặt trùng.")
add_normal("Bước 1 - Chọn sân và khung giờ: Người dùng chọn sân, chọn ngày và xem các khung giờ trống. Giao diện cho phép chọn nhiều slot liên tiếp.")
add_normal("Bước 2 - Khóa chỗ (Slot Lock): Khi người dùng chọn các khung giờ, hệ thống tạo các bản ghi khóa chỗ (SlotLock) trong cơ sở dữ liệu với thời hạn 10-15 phút. Trong thời gian này, các slot đó không thể được đặt bởi người dùng khác. Mỗi slot có trạng thái LOCKED và hạn expires_at để tự động mở khóa.")
add_normal("Bước 3 - Tạo đơn đặt sân: Người dùng xác nhận đặt sân. Hệ thống kiểm tra các slot còn trong trạng thái LOCKED, tạo đơn đặt sân mới với trạng thái PENDING, chuyển các slot sang trạng thái CONVERTED_TO_BOOKING và phát sự kiện BookingCreated.")
add_normal("Bước 4 - Thanh toán: Người dùng được chuyển đến trang thanh toán để hoàn tất giao dịch.")
add_normal("Bước 5 - Xác nhận: Sau khi thanh toán thành công, đơn đặt sân chuyển sang trạng thái PAID. Người dùng và chủ sân nhận thông báo xác nhận.")
add_normal("Cơ chế tự động hết hạn: Một công việc định kỳ (scheduled job) chạy trong nền để kiểm tra và hết hạn các đơn đặt sân PENDING và slot LOCKED đã quá thời gian cho phép, tự động chuyển sang trạng thái EXPIRED.")

add_heading_custom("Usecase thanh toán đơn đặt sân", 3)
add_normal("Chức năng thanh toán cho phép người dùng thanh toán cho đơn đặt sân thông qua cổng thanh toán tích hợp.")
add_normal("Luồng thanh toán: Sau khi tạo đơn đặt sân, hệ thống tạo một giao dịch thanh toán với trạng thái PENDING. Người dùng được chuyển đến trang thanh toán (mock hoặc VNPay). Sau khi hoàn tất thanh toán, cổng thanh toán gọi callback về hệ thống. Payment Service cập nhật trạng thái giao dịch và phát sự kiện PaymentSucceeded hoặc PaymentFailed. Booking Service tiêu thụ sự kiện này để cập nhật trạng thái đơn đặt sân tương ứng.")
add_normal("Phương thức thanh toán: Hiện tại hệ thống hỗ trợ thanh toán thử nghiệm (mock) và đang phát triển tích hợp VNPay.")

add_heading_custom("Usecase đánh giá sân", 3)
add_normal("Chức năng đánh giá sân cho phép người chơi đánh giá chất lượng sân sau khi đã sử dụng dịch vụ.")
add_normal("Điều kiện đánh giá: Người dùng chỉ có thể đánh giá sân nếu đã có đơn đặt sân tại sân đó với trạng thái PAID (đã thanh toán thành công). Điều kiện này được kiểm tra thông qua internal API call từ Venue Service đến Booking Service.")
add_normal("Nội dung đánh giá: Người dùng chọn số sao (1-5) và có thể thêm nhận xét về sân. Đánh giá có thể được cập nhật (không phải đánh giá một lần).")
add_normal("Cập nhật điểm trung bình: Mỗi khi có đánh giá mới hoặc cập nhật, Venue Service tự động tính toán lại điểm trung bình (rating_avg) và tổng số đánh giá (rating_count) của sân.")

add_heading_custom("Phân tích chi tiết chức năng cộng đồng (Community)", 2)

add_heading_custom("Usecase tạo bài đăng kèo giao lưu", 3)
add_normal("Chức năng tạo bài đăng kèo giao lưu cho phép người chơi tạo các trận giao lưu để tìm đối thủ chơi cầu lông cùng.")
add_normal("Luồng nghiệp vụ: Người chơi nhập thông tin kèo bao gồm tiêu đề, mô tả, địa điểm (chọn sân có sẵn hoặc nhập địa chỉ tùy chỉnh), thời gian bắt đầu và kết thúc, trình độ yêu cầu, số người tối đa và chế độ tham gia (mở hoặc cần phê duyệt). Hệ thống kiểm tra tính hợp lệ của dữ liệu: tiêu đề từ 3 đến 120 ký tự, thời gian bắt đầu ít nhất 30 phút từ hiện tại, thời gian kết thúc sau thời gian bắt đầu, số người tối đa ít nhất 2 người.")
add_normal("Chế độ tham gia: Chế độ OPEN cho phép người chơi tham gia ngay lập tức mà không cần chờ phê duyệt. Chế độ APPROVAL yêu cầu chủ kèo phê duyệt từng người tham gia.")
add_normal("Phát sự kiện: Sau khi tạo bài đăng thành công, Community Service phát sự kiện MatchPostCreated đến Kafka, Notification Service tiêu thụ để gửi thông báo cho các người chơi phù hợp.")

add_heading_custom("Usecase tham gia kèo giao lưu", 3)
add_normal("Chức năng tham gia kèo cho phép người chơi gửi yêu cầu tham gia các kèo giao lưu đang mở.")
add_normal("Luồng nghiệp vụ: Người chơi chọn một kèo giao lưu đang mở và gửi yêu cầu tham gia. Nếu kèo ở chế độ OPEN, người chơi được thêm vào danh sách tham gia ngay lập tức (trạng thái APPROVED) và Notification Service gửi thông báo cho chủ kèo. Nếu kèo ở chế độ APPROVAL, yêu cầu được lưu với trạng thái PENDING, Notification Service gửi thông báo cho chủ kèo để phê duyệt.")
add_normal("Ràng buộc: Người dùng không thể tham gia kèo đã đủ số người tối đa, kèo đã đóng/tạm đóng, hoặc kèo mà mình đã tham gia (cơ chế idempotent).")
add_normal("Tự động đóng kèo: Khi số người tham gia đã được phê duyệt đạt số người tối đa, kèo tự động chuyển sang trạng thái CLOSED.")

add_heading_custom("Usecase quản lý người tham gia", 3)
add_normal("Chức năng quản lý người tham gia cho phép chủ kèo phê duyệt hoặc từ chối các yêu cầu tham gia vào kèo của mình.")
add_normal("Phê duyệt/Từ chối: Chủ kèo xem danh sách người chơi chờ phê duyệt, có thể phê duyệt hoặc từ chối từng người với lý do. Khi phê duyệt, trạng thái người chơi chuyển sang APPROVED và sự kiện MatchApproved được phát. Khi từ chối, trạng thái chuyển sang REJECTED.")
add_normal("Rời khỏi kèo: Người chơi có thể rời khỏi kèo mình đã tham gia (trạng thái chuyển sang CANCELLED_BY_USER). Chủ kèo có thể xóa người chơi khỏi kèo (trạng thái chuyển sang REMOVED_BY_HOST).")
add_normal("Kết thúc kèo: Chủ kèo có thể kết thúc kèo (trạng thái chuyển sang FINISHED) sau khi trận đấu kết thúc, mở khóa chức năng đánh giá người chơi.")

add_heading_custom("Usecase đánh giá người chơi", 3)
add_normal("Chức năng đánh giá người chơi cho phép chủ kèo đánh giá chất lượng chơi của các người tham gia sau khi trận đấu kết thúc.")
add_normal("Điều kiện đánh giá: Chỉ chủ kèo mới có quyền đánh giá, kèo phải ở trạng thái FINISHED, người được đánh giá phải có trạng thái APPROVED trong danh sách tham gia.")
add_normal("Nội dung đánh giá: Chủ kèo chọn số sao (1-5) và có thể thêm nhận xét về người chơi. Đánh giá có thể được cập nhật.")
add_normal("Phát sự kiện: Sau khi đánh giá, sự kiện RatingCreated được phát đến Kafka, Notification Service thông báo cho người được đánh giá.")

add_heading_custom("Thiết kế cơ sở dữ liệu", 2)

add_heading_custom("Mô hình dữ liệu", 3)
add_normal("Hệ thống sử dụng chiến lược đa cơ sở dữ liệu (polyglot persistence), trong đó mỗi loại database được lựa chọn phù hợp với đặc thù nghiệp vụ:")
add_normal("PostgreSQL (đa schema): Là cơ sở dữ liệu chính, lưu trữ toàn bộ dữ liệu nghiệp vụ cốt lõi. Sử dụng cơ chế đa schema để tách biệt dữ liệu giữa các microservices: schema identity (người dùng và xác thực), schema venue (sân và thông tin sân), schema booking (đặt sân và khóa chỗ), schema payment (thanh toán) và schema community (cộng đồng và kèo giao lưu).")
add_normal("MongoDB: Lưu trữ dữ liệu phi cấu trúc cho module thông báo và gợi ý với hai database riêng biệt: notification_db (thông báo) và recommendation_db (hoạt động người dùng).")
add_normal("Redis: Lưu trữ dữ liệu volatile với thời gian sống ngắn: OTP khôi phục mật khẩu, bộ nhớ đệm gợi ý và quản lý khóa chỗ sân.")

add_heading_custom("Chi tiết các thực thể chính", 3)
add_normal("Các thực thể chính trong hệ thống được thiết kế như sau:")
add_normal("User (identity schema): Lưu trữ thông tin tài khoản người dùng bao gồm id (UUID), email, số điện thoại, mật khẩu (đã mã hóa BCrypt), họ tên, trạng thái tài khoản (ACTIVE/LOCKED) và thời gian tạo.")
add_normal("Role: Lưu trữ thông tin vai trò trong hệ thống với mã vai trò (USER/OWNER/ADMIN).")
add_normal("RefreshToken: Lưu trữ refresh token của người dùng với token_hash (đã mã hóa), thời gian hết hạn và trạng thái thu hồi.")
add_normal("UserProfile: Lưu trữ thông tin mở rộng của người dùng bao gồm trình độ chơi, giới tính, mục tiêu, khu vực ưa thích, tiểu sử, điểm đánh giá trung bình và tổng số đánh giá.")
add_normal("Venue (venue schema): Lưu trữ thông tin sân cầu lông bao gồm id, owner_id (chủ sân), tên sân, địa chỉ, thành phố, quận/huyện, tọa độ địa lý (geometry - PostGIS), trạng thái (PENDING_APPROVAL/APPROVED/REJECTED), điểm đánh giá trung bình, tổng số đánh giá và tiện ích.")
add_normal("Court: Lưu trữ thông tin sân con bao gồm id, venue_id, tên sân con, loại sân (đơn/đôi) và trạng thái hoạt động.")
add_normal("CourtSlot: Lưu trữ khung giờ hoạt động định kỳ của sân con theo ngày trong tuần, bao gồm ngày, giờ bắt đầu, giờ kết thúc và trạng thái.")
add_normal("PriceRule: Lưu trữ quy tắc giá linh hoạt, áp dụng theo loại ngày (thường/cuối tuần), khoảng thời gian trong ngày và sân con cụ thể.")
add_normal("VenueRating: Lưu trữ đánh giá sân của người dùng với số sao (1-5), nhận xét và liên kết đến đơn đặt sân đã thanh toán để xác minh.")
add_normal("Booking (booking schema): Lưu trữ thông tin đơn đặt sân với user_id, venue_id, court_id, tổng tiền, trạng thái (PENDING/PAID/FAILED/EXPIRED/CANCELLED_BY_ADMIN), trạng thái thanh toán và thời gian tạo.")
add_normal("BookingItem: Lưu trữ chi tiết từng khung giờ trong đơn đặt sân, bao gồm sân con, thời gian bắt đầu, kết thúc, giá snapshot và trạng thái.")
add_normal("SlotLock: Lưu trữ khóa chỗ sân với user_id, venue_id, court_id, thời gian bắt đầu, kết thúc, trạng thái (LOCKED/RELEASED/EXPIRED/CONVERTED_TO_BOOKING) và thời gian hết hạn.")
add_normal("PaymentTransaction (payment schema): Lưu trữ giao dịch thanh toán với booking_id, user_id, số tiền, nhà cung cấp thanh toán (MOCK/VNPAY), trạng thái, URL thanh toán, mã giao dịch và thời gian tạo.")
add_normal("MatchPost (community schema): Lưu trữ bài đăng kèo giao lưu với author_id, tiêu đề, mô tả, địa điểm (venue_id hoặc location_text + tọa độ), thời gian, trình độ yêu cầu, số người tối đa, chế độ tham gia (OPEN/APPROVAL), trạng thái (OPEN/CLOSED/FINISHED/HIDDEN).")
add_normal("Participant: Lưu trữ người tham gia kèo với match_post_id, user_id, tên người dùng, trạng thái (PENDING/APPROVED/REJECTED/CANCELLED_BY_USER/REMOVED_BY_HOST), ghi chú và thời gian tham gia.")
add_normal("PlayerRating: Lưu trữ đánh giá người chơi với match_post_id, người đánh giá, người được đánh giá, số sao và nhận xét.")

add_heading_custom("Thiết kế giao diện người dùng", 2)

add_heading_custom("Giao diện người dùng (User)", 3)
add_normal("Giao diện dành cho người chơi được thiết kế tập trung vào trải nghiệm tìm kiếm và đặt sân một cách nhanh chóng, thân thiện với người dùng.")
add_normal("Trang chủ: Hiển thị thông tin giới thiệu về nền tảng, các tính năng nổi bật và lời kêu gọi đăng ký sử dụng. Người dùng đã đăng nhập có thể nhanh chóng truy cập các chức năng chính như tìm sân, tạo kèo.")
add_normal("Danh sách sân: Hiển thị các sân công khai dưới dạng thẻ (card) với hình ảnh, tên sân, địa chỉ, giá từ và đánh giá trung bình. Người dùng có thể lọc theo thành phố, quận/huyện, khoảng giá và tiện ích. Hỗ trợ chuyển đổi giữa chế độ danh sách và bản đồ.")
add_normal("Chi tiết sân: Hiển thị thông tin đầy đủ về sân bao gồm mô tả, hình ảnh, tiện ích, danh sách sân con, khung giờ hoạt động, bảng giá và đánh giá từ người chơi khác. Giao diện cho phép chọn ngày và xem các khung giờ trống để đặt sân.")
add_normal("Đặt sân: Giao diện chọn khung giờ với khả năng kéo thả (drag & drop) để chọn nhiều slot liên tiếp. Hiển thị tổng tiền dự kiến và nút xác nhận đặt sân.")
add_normal("Thanh toán: Trang thanh toán hiển thị thông tin đơn đặt sân và các phương thức thanh toán. Người dùng xác nhận thanh toán và nhận kết quả.")
add_normal("Lịch sử đặt sân: Danh sách các đơn đặt sân đã thực hiện, có thể lọc theo trạng thái và theo thời gian.")
add_normal("Cộng đồng: Hiển thị danh sách bài đăng kèo giao lưu, cho phép lọc theo khu vực, trình độ và thời gian. Người dùng có thể tạo kèo mới hoặc tham gia kèo có sẵn.")
add_normal("Chi tiết kèo: Hiển thị thông tin kèo, danh sách người tham gia và cho phép tham gia hoặc rời khỏi kèo.")
add_normal("Thông báo: Trung tâm thông báo hiển thị tất cả thông báo về đơn đặt sân, lời mời tham gia kèo, kết quả đánh giá và các hoạt động liên quan.")
add_normal("Hồ sơ cá nhân: Trang quản lý thông tin cá nhân bao gồm tên, giới tính, trình độ chơi, mục tiêu và khu vực ưa thích.")

add_heading_custom("Giao diện chủ sân (Owner)", 3)
add_normal("Giao diện dành cho chủ sân được thiết kế theo phong cách dashboard với các biểu đồ và bảng dữ liệu, tập trung vào việc quản lý hiệu quả cơ sở kinh doanh.")
add_normal("Dashboard: Hiển thị tổng quan doanh thu, số đơn đặt sân mới, công suất sân và các chỉ số quan trọng. Các biểu đồ trực quan hóa xu hướng doanh thu và lịch sử đặt sân.")
add_normal("Quản lý sân: Danh sách các sân thuộc quản lý, nút thêm sân mới. Xem chi tiết và chỉnh sửa thông tin từng sân.")
add_normal("Quản lý sân con: Thêm, sửa, xóa các sân con. Cấu hình loại sân và trạng thái hoạt động.")
add_normal("Quản lý khung giờ và giá: Giao diện trực quan để thiết lập khung giờ hoạt động và các quy tắc giá theo thời gian.")
add_normal("Quản lý đơn đặt sân: Danh sách các đơn đặt sân đến với khả năng lọc theo trạng thái và thời gian. Xem chi tiết đơn đặt và theo dõi trạng thái thanh toán.")
add_normal("Doanh thu: Trang thống kê doanh thu với biểu đồ phân tích theo ngày, tuần, tháng. Xem doanh thu chi tiết theo từng sân con.")

add_heading_custom("Giao diện quản trị viên (Admin)", 3)
add_normal("Giao diện quản trị viên được thiết kế theo phong cách dashboard chuyên nghiệp với các chỉ số tổng quan và công cụ quản lý toàn hệ thống.")
add_normal("Dashboard: Hiển thị tổng quan hệ thống bao gồm số lượng người dùng, số sân đang hoạt động, số đơn đặt sân trong ngày và doanh thu tổng. Biểu đồ xu hướng đăng ký người dùng và đơn đặt sân.")
add_normal("Quản lý người dùng: Bảng danh sách tất cả người dùng với khả năng lọc theo vai trò và trạng thái. Tìm kiếm theo email hoặc tên. Hành động khóa/mở khóa tài khoản, thay đổi vai trò và xóa người dùng được thực hiện trực tiếp trên bảng.")
add_normal("Phê duyệt sân: Danh sách sân chờ phê duyệt với thông tin chi tiết. Xem xem thông tin và quyết định phê duyệt hoặc từ chối (với lý do).")
add_normal("Quản lý đơn đặt sân: Xem tất cả đơn đặt sân trên hệ thống, có thể hủy đơn với lý do.")
add_normal("Báo cáo: Trang báo cáo tổng quan với các biểu đồ thống kê hoạt động hệ thống.")

add_heading_custom("Giao diện khách vãng lai", 3)
add_normal("Giao diện dành cho khách vãng lai (chưa đăng nhập) được thiết kế đơn giản, tập trung vào việc giới thiệu nền tảng và khuyến khích đăng ký.")
add_normal("Trang chủ công khai: Giới thiệu các tính năng chính của hệ thống, lời kêu gọi đăng ký và đăng nhập.")
add_normal("Xem sân công khai: Khách có thể duyệt danh sách sân đã được phê duyệt và xem thông tin chi tiết sân (không thể đặt sân).")
add_normal("Đăng ký/Đăng nhập: Các nút đăng ký và đăng nhập luôn có sẵn để khách dễ dàng tham gia hệ thống.")

doc.add_page_break()

# =============================
# CHƯƠNG 5: XÂY DỰNG HỆ THỐNG & KIỂM THỬ
# =============================
add_heading_custom("XÂY DỰNG HỆ THỐNG & KIỂM THỬ", 1)

add_heading_custom("Xây dựng hệ thống", 2)

add_heading_custom("Kiến trúc triển khai", 3)
add_normal("Hệ thống được triển khai theo kiến trúc microservices với sử dụng Docker và Docker Compose để quản lý toàn bộ hạ tầng. Kiến trúc này mang lại nhiều ưu điểm:")
add_normal("Đóng gói ứng dụng: Mỗi microservice được đóng gói thành Docker image với multi-stage build, đảm bảo image cuối cùng nhẹ và chỉ chứa những gì cần thiết để chạy.")
add_normal("Orchestration: Docker Compose định nghĩa và quản lý toàn bộ stack gồm 10+ microservices, cơ sở dữ liệu (PostgreSQL, MongoDB, Redis), message broker (Kafka) và các công cụ hỗ trợ (Eureka, Config Server, Kafka UI).")
add_normal("Môi trường đồng nhất: Docker đảm bảo môi trường chạy nhất quán giữa máy phát triển và máy triển khai, loại bỏ vấn đề \"chạy được trên máy tôi\". Hệ thống có các script khởi động và dừng dịch vụ (start-services.bat/sh và stop-services.bat/sh) để dễ dàng vận hành.")
add_normal("Service Discovery: Netflix Eureka đóng vai trò registry trung tâm, các microservice tự động đăng ký khi khởi động và có thể tìm kiếm lẫn nhau mà không cần hardcode địa chỉ.")

add_heading_custom("Công nghệ triển khai", 3)
add_normal("Backend - Microservices:")
add_normal("Mỗi microservice được xây dựng như một ứng dụng Spring Boot độc lập, có thể được build, test và triển khai riêng lẻ. Các dịch vụ giao tiếp thông qua REST API (qua API Gateway hoặc Feign client) và sự kiện Kafka (không đồng bộ).")
add_normal("Mỗi dịch vụ có cấu trúc code theo pattern chuẩn: Controller xử lý HTTP request, Service chứa logic nghiệp vụ, Repository truy cập dữ liệu, DTO chuyển đổi dữ liệu, Exception xử lý lỗi.")
add_normal("Frontend - Single Page Application:")
add_normal("Ứng dụng React được build bằng Vite, tạo ra các static assets được phục vụ bởi Nginx trong môi trường production. Router quản lý điều hướng với các route guard dựa trên vai trò người dùng.")
add_normal("Axios interceptor tự động gắn access token vào mỗi request và xử lý refresh token khi access token hết hạn, đảm bảo trải nghiệm người dùng liền mạch.")
add_normal("Infrastructure:")
add_normal("PostgreSQL với PostGIS cung cấp cơ sở dữ liệu chính, được khởi tạo với các schema và dữ liệu mẫu thông qua script SQL.")
add_normal("MongoDB lưu trữ dữ liệu thông báo và gợi ý, được khởi tạo với các collection và index phù hợp.")
add_normal("Redis cung cấp bộ nhớ đệm và quản lý OTP, được cấu hình với chính sách TTL phù hợp.")
add_normal("Apache Kafka đóng vai trò event bus trung tâm với 4 topics (partner-onboarding, payment.events, booking.events, community.events) và 8 loại sự kiện khác nhau.")
add_normal("Kafka UI được tích hợp để theo dõi và quản lý các sự kiện trong hệ thống, hỗ trợ việc debug và giám sát luồng sự kiện giữa các microservices.")
add_normal("Nginx đóng vai trò reverse proxy cho API Gateway và web server cho Frontend, xử lý SSL termination và static file serving.")

add_heading_custom("Kết quả xây dựng", 3)
add_normal("Hệ thống được xây dựng hoàn chỉnh với đầy đủ các chức năng cốt lõi theo đặc tả ban đầu:")
add_normal("Backend - 10 microservices đã được phát triển và tích hợp thành công:")
add_normal("API Gateway (8080): Xử lý routing, xác thực JWT và áp dụng các chính sách bảo mật.")
add_normal("Identity Service (8081): Xác thực người dùng, phân quyền, quản lý hồ sơ và khôi phục mật khẩu OTP.")
add_normal("Venue Service (8082): Quản lý sân, sân con, giá cả, đánh giá sân và tìm kiếm theo vị trí PostGIS.")
add_normal("Booking Service (8083): Quản lý đặt sân với cơ chế khóa chỗ tự động và xử lý hết hạn.")
add_normal("Payment Service (8084): Xử lý thanh toán với tích hợp cổng thanh toán.")
add_normal("Community Service (8085): Quản lý kèo giao lưu, tham gia kèo, đánh giá người chơi và tìm kiếm theo vị trí.")
add_normal("Notification Service (8086): Gửi thông báo cho người dùng dựa trên sự kiện từ các microservice khác.")
add_normal("Recommendation Service (8087): Cung cấp gợi ý sân và kèo thông minh với bộ nhớ đệm Redis.")
add_normal("Chat Service (8686): Hỗ trợ nhắn tin thời gian thực giữa người chơi.")
add_normal("AI Service (8091): Tích hợp trí tuệ nhân tạo cho các tính năng thông minh.")
add_normal("Frontend - 50+ trang và thành phần được xây dựng hoàn chỉnh, bao gồm các trang dành cho người chơi, chủ sân, quản trị viên và khách vãng lai. Giao diện responsive, hỗ trợ cả máy tính để bàn và thiết bị di động.")
add_normal("Cơ sở dữ liệu - 5 schema PostgreSQL, 2 database MongoDB và Redis được thiết kế và khởi tạo với cấu trúc tối ưu cho từng loại nghiệp vụ.")
add_normal("Hệ thống sự kiện - 4 Kafka topics với 8+ loại sự kiện, đảm bảo giao tiếp không đồng bộ giữa các microservices theo mô hình event-driven.")
add_normal("Hệ thống cung cấp tài liệu API tự động (Swagger UI) cho mỗi microservice và bộ sưu tập Postman để kiểm thử API.")

add_heading_custom("Kiểm thử", 2)

add_heading_custom("Phương pháp kiểm thử", 3)
add_normal("Hệ thống được kiểm thử thông qua nhiều phương pháp khác nhau để đảm bảo chất lượng:")
add_normal("Kiểm thử API: Mỗi microservice được trang bị Swagger UI, cho phép thực hiện kiểm thử API trực tiếp thông qua giao diện web. Bộ sưu tập Postman được xây dựng để kiểm thử toàn bộ các endpoint của hệ thống một cách có hệ thống.")
add_normal("Kiểm thử tích hợp: Kiểm thử luồng nghiệp vụ xuyên suốt các microservice, bao gồm luồng đăng ký đăng nhập, luồng đặt sân từ lúc khóa chỗ đến thanh toán, luồng tạo và tham gia kèo giao lưu.")
add_normal("Kiểm thử chức năng: Kiểm thử từng chức năng riêng lẻ để đảm bảo hoạt động đúng theo đặc tả, bao gồm đăng ký/đăng nhập, CRUD sân, đặt sân, thanh toán, tạo kèo và các chức năng quản lý.")
add_normal("Kiểm thử cơ chế khóa chỗ: Kiểm tra cơ chế khóa chỗ sân hoạt động đúng, ngăn chặn đặt trùng và tự động hết hạn khóa chỗ sau thời gian quy định.")

add_heading_custom("Kết quả kiểm thử", 3)
add_normal("Các kết quả kiểm thử chính sau khi triển khai hệ thống:")
add_normal("Luồng đặt sân: Chức năng đặt sân hoạt động đúng từ bước chọn khung giờ, khóa chỗ, tạo đơn đặt sân đến thanh toán. Cơ chế khóa chỗ ngăn chặn thành công tình trạng đặt trùng khi nhiều người dùng cùng đặt cùng một slot.")
add_normal("Luồng đăng ký và đăng nhập: Xác thực người dùng hoạt động ổn định với JWT access token và refresh token. Phân quyền theo vai trò được áp dụng chính xác trên tất cả các endpoint.")
add_normal("Luồng cộng đồng: Tạo bài đăng kèo, tham gia kèo (cả chế độ OPEN và APPROVAL), phê duyệt người tham gia và đánh giá người chơi đều hoạt động đúng theo quy tắc nghiệp vụ đã định nghĩa.")
add_normal("Luồng quản lý: Chủ sân có thể đăng ký sân mới (chờ phê duyệt), quản lý thông tin sân, xem đơn đặt sân và doanh thu. Quản trị viên có thể phê duyệt sân, quản lý người dùng và xem báo cáo tổng quan.")
add_normal("Giao tiếp microservices: Các microservice giao tiếp đúng thông qua API Gateway và Kafka. Sự kiện được phát và tiêu thụ chính xác giữa các dịch vụ (BookingCreated, PaymentSucceeded, MatchPostCreated, v.v.).")
add_normal("Tài khoản thử nghiệm: Hệ thống được kiểm thử với các tài khoản mẫu sau: Quản trị viên (admin@badmintonhub.local), Chủ sân (owner@badminton.com), Người chơi (user@badminton.com). Mỗi tài khoản được kiểm tra đầy đủ các chức năng tương ứng với vai trò.")

doc.add_page_break()

# =============================
# CHƯƠNG 6: KẾT LUẬN
# =============================
add_heading_custom("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN TƯƠNG LAI", 1)

add_heading_custom("Kết luận", 2)
add_normal("Sau thời gian thực hiện đồ án, với những kiến thức được học và tích lũy của bản thân, em đã hoàn thành đồ án tốt nghiệp đáp ứng đầy đủ các yêu cầu và mục tiêu đã đề ra:")
add_normal("Nghiên cứu phương pháp và các vấn đề lý thuyết liên quan đến kiến trúc microservices, hệ thống giao tiếp sự kiện không đồng bộ (event-driven), phát triển ứng dụng web với React và Spring Boot, tích hợp cơ sở dữ liệu đa dạng (PostgreSQL/PostGIS, MongoDB, Redis) và hệ thống message broker Apache Kafka.")
add_normal("Khảo sát thực trạng về hoạt động đặt sân cầu lông và nhu cầu kết nối cộng đồng người chơi, từ đó phân tích và thiết kế quy trình nghiệp vụ toàn diện cho hệ thống, đảm bảo giải quyết được các bất cập trong phương thức đặt sân truyền thống.")
add_normal("Phân tích và thiết kế kiến trúc tổng thể hệ thống theo mô hình microservices với API Gateway, Service Registry, event-driven communication qua Kafka, và cơ sở dữ liệu đa schema. Thiết kế chi tiết từng module với usecase, biểu đồ lớp và cấu trúc cơ sở dữ liệu.")
add_normal("Phát triển thành công hệ thống đặt sân cầu lông trực tuyến tích hợp cộng đồng giao lưu và gợi ý thông minh với các tính năng cốt lõi: đặt sân trực tuyến với cơ chế khóa chỗ tự động, thanh toán tích hợp, tìm kiếm sân theo vị trí, cộng đồng kèo giao lưu, đánh giá sân và người chơi, thông báo thời gian thực và gợi sân/kèo thông minh.")
add_normal("Triển khai toàn bộ hệ thống trên nền tảng Docker với Docker Compose, đảm bảo khả năng triển khai và vận hành dễ dàng trên nhiều môi trường khác nhau.")

add_heading_custom("Hướng phát triển tương lai", 2)
add_normal("Mặc dù đã hoàn thành các tính năng cốt lõi và đạt được các mục tiêu đề ra, hệ thống vẫn còn nhiều hướng phát triển trong tương lai để trở nên hoàn thiện và sẵn sàng đưa vào vận hành thương mại:")
add_normal("Tích hợp cổng thanh toán thực tế: Hoàn thiện tích hợp VNPay để thay thế phương thức thanh toán thử nghiệm (mock), cho phép người dùng thực hiện giao dịch thanh toán thực tế với đa dạng phương thức (thẻ, ví điện tử, chuyển khoản).")
add_normal("Hệ thống email thông báo: Tích hợp dịch vụ gửi email (SendGrid hoặc AWS SES) để gửi OTP khôi phục mật khẩu và các thông báo quan trọng đến hộp thư của người dùng thay vì hiển thị trên console.")
add_normal("Tính năng nhắn tin thời gian thực: Phát triển đầy đủ tính năng chat thời gian thực giữa người chơi thông qua WebSocket, bao gồm lịch sử tin nhắn, thông báo tin nhắn mới và trạng thái online/offline.")
add_normal("Cơ chế đồng bộ sự kiện đáng tin cậy: Triển khai mô hình Outbox Pattern và cơ chế xử lý idempotent cho Kafka events để đảm bảo không có sự kiện nào bị mất hoặc xử lý trùng lặp trong môi trường production.")
add_normal("Giới hạn tần suất yêu cầu (Rate Limiting): Triển khai cơ chế giới hạn số lượng yêu cầu API trong một khoảng thời gian nhất định bằng Redis để bảo vệ hệ thống khỏi các cuộc tấn công DDoS và sử dụng quá mức.")
add_normal("Tính năng bình luận và thích bài đăng: Phát triển API cho tính năng bình luận (Comment), thích (Like) và báo cáo (Report) bài đăng trong cộng đồng kèo giao lưu, hiện tại các thực thể đã được thiết kế nhưng chưa có API triển khai.")
add_normal("Đề xuất người chơi thông minh: Phát triển thuật toán gợi ý người chơi phù hợp dựa trên trình độ, khu vực và lịch sử tham gia kèo, bổ sung cho chức năng gợi ý sân đã có sẵn.")
add_normal("Kiểm thử tự động hóa: Xây dựng bộ kiểm thử tự động (unit test, integration test, E2E test) với các framework phù hợp để đảm bảo chất lượng code và ngăn ngừa regression khi phát triển thêm tính năng mới.")
add_normal("Giám sát và quan sát hệ thống: Tích hợp Prometheus và Grafana để giám sát hiệu năng các microservices, theo dõi tình trạng sức khỏe hệ thống và cảnh báo sự cố kịp thời.")
add_normal("Xây dựng tài liệu API chi tiết: Bổ sung ví dụ request/response chi tiết vào Swagger UI của từng microservice, xây dựng tài liệu hướng dẫn tích hợp cho bên thứ ba.")

# Save
doc.save('D:/Hoc_ki_8/DoAn/BaoCaoDoAn_NenTangDatSanCauLong.docx')
print("DONE - File saved successfully!")
